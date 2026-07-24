from flask import Flask, render_template, jsonify, request, send_file
import json, os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

BASE = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE, "data.json")

app = Flask(__name__)
app.config["SEND_FILE_MAX_AGE_DEFAULT"] = 0

@app.context_processor
def inject_globals():
    theme = None
    try:
        theme = request.cookies.get("theme", "light")
    except Exception:
        theme = "light"
    return {"now": datetime.now(), "request_theme": theme}

@app.route("/api/theme", methods=["POST"])
def set_theme():
    theme = request.json.get("theme") if request.json else None
    if theme not in ("light", "dark"):
        theme = "light"
    resp = jsonify({"ok": True, "theme": theme})
    resp.set_cookie("theme", theme, max_age=60 * 60 * 24 * 365, secure=False, httponly=False, samesite="Lax")
    return resp

@app.route("/")
def dashboard():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("dashboard.html", data=data, theme=theme)

@app.route("/about")
def about():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("about.html", data=data, theme=theme)

@app.route("/profile")
def profile():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("profile.html", data=data, theme=theme)

@app.route("/skills")
def skills():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("skills.html", data=data, theme=theme)

@app.route("/experience")
def experience():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("experience.html", data=data, theme=theme)

@app.route("/projects")
def projects():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("projects.html", data=data, theme=theme)

@app.route("/resume")
def resume_page():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("resume.html", data=data, theme=theme)

@app.route("/certifications")
def certifications():
    data = load_data()
    theme = request.cookies.get("theme", "light")
    return render_template("certifications.html", data=data, theme=theme)

@app.route("/api/resume/docx")
def resume_docx():
    data = load_data()
    p = data.get("profile", {})
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Helper: section divider + centered uppercase heading + thin top border
    def add_section_heading(doc, text):
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Calibri"
        ppr = header._element.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="1" w:color="000000"/></w:pBdr>')
        ppr.append(pBdr)

    def add_section_body(doc):
        # reduced spacing; no extra blank paragraph
        pass

    # Header
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(p.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(26)
    name_run.font.name = "Calibri"

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = " | ".join([v for v in [p.get("phone",""), p.get("email",""), "LinkedIn", p.get("location","")] if v])
    contact.add_run(contact_text)

    # thick bottom border for header
    ppr = name_para._element.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/></w:pBdr>')
    ppr.append(pBdr)
    doc.add_paragraph()  # spacing

    # Profile
    add_section_heading(doc, "PROFILE")
    doc.add_paragraph(p.get("summary", ""))
    add_section_body(doc)

    # Professional Experience
    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
    for e in data.get("experience", []):
        # Job header row: title/location left, date right, no table borders
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = True
        tbl.allow_autofit = True
        tbl.columns[0].width = Inches(5.5)
        tbl.columns[1].width = Inches(1.5)
        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]
        left.text = f"{e.get('role','')}, {e.get('company','')}, {e.get('location','')}"
        right.text = f"{e.get('start','')} – {e.get('end') or 'Present'}"
        # Make left bold, right bold
        for cell in (left, right):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if cell == left else WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
        # remove borders
        tbl.style = None
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>')
                tcPr.append(tcBorders)
        for a in e.get("achievements", []):
            p = doc.add_paragraph(a, style="List Bullet")
        if e.get("tech"):
            doc.add_paragraph("Tech: " + ", ".join(e["tech"]))
    add_section_body(doc)

    # Education
    add_section_heading(doc, "EDUCATION")
    edus = data.get("education", [])
    for e in edus:
        if not e.get("school"):
            continue
        doc.add_paragraph(e.get("degree", ""), style="Normal")
        doc.add_paragraph(f"{e.get('start','')} – {e.get('end','')}")
        doc.add_paragraph(e.get("school", ""))
    add_section_body(doc)

    # Certifications
    add_section_heading(doc, "CERTIFICATIONS")
    certs = data.get("certifications", [])
    if certs:
        for c in certs:
            doc.add_paragraph(f"{c.get('name','')} — {c.get('issuer','')}", style="List Bullet")
    else:
        doc.add_paragraph("No certifications listed yet.", style="Normal")
    add_section_body(doc)

    # Key Skills
    add_section_heading(doc, "KEY SKILLS")
    skills = data.get("skills", [])
    mid = (len(skills) + 1) // 2
    left_skills = skills[:mid]
    right_skills = skills[mid:]
    table = doc.add_table(rows=max(len(left_skills), len(right_skills), 1), cols=2)
    table.style = None
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    for idx in range(len(left_skills)):
        left_cell = table.rows[idx].cells[0]
        left_cell.text = "• " + left_skills[idx].get("name", "")
    for idx in range(len(right_skills)):
        right_cell = table.rows[idx].cells[1]
        right_cell.text = "• " + right_skills[idx].get("name", "")
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>')
            tcPr.append(tcBorders)
    add_section_body(doc)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="resume.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def load_data():
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_data(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

if __name__ == "__main__":
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", 5678))
    debug = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(host=host, port=port, debug=debug)
